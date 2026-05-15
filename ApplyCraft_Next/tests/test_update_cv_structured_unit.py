import os
import re
import shutil
import tempfile
import unittest

from docx import Document

from core.config import JOB_POSITIONS
from core.update_cv import update_cv_bullets


def _split_legacy_job_title(job_title):
    parts = re.split(r"\s+[^\w\s]+\s+", job_title, maxsplit=1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    return job_title, ""


class TestStructuredUpdateCV(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.mkdtemp(prefix="applycraft_update_cv_test_")
        self.template_path = os.path.join(self.temp_dir, "template.docx")
        self.output_path = os.path.join(self.temp_dir, "output.docx")

    def tearDown(self):
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def _build_template(self):
        legacy_key = list(JOB_POSITIONS.keys())[0]
        company, title = _split_legacy_job_title(legacy_key)

        doc = Document()
        doc.add_paragraph("Old Name")
        doc.add_paragraph("old@email.com")
        doc.add_paragraph("old-link")
        doc.add_paragraph("Summary")
        doc.add_paragraph("Old summary text.")
        doc.add_paragraph("Work Experience")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = company
        table.cell(0, 1).text = "Old Location"
        table.cell(1, 0).text = title
        table.cell(1, 1).text = "Old Dates"
        doc.add_paragraph("- old bullet 1", style="List Bullet")
        doc.add_paragraph("- old bullet 2", style="List Bullet")
        doc.add_paragraph("Skills")
        doc.save(self.template_path)
        return legacy_key

    def test_structured_experience_updates_headers_and_bullets(self):
        legacy_key = self._build_template()
        experiences = [
            {
                "anchor_key": "exp_1",
                "legacy_key": legacy_key,
                "company": "Manual Company",
                "title": "Manual Role",
                "date_range": "2024 - Present",
                "location": "Berlin, Germany",
                "headline": "Impact headline",
                "bullets": ["New bullet one"],
            }
        ]

        update_cv_bullets(
            input_file=self.template_path,
            output_file=self.output_path,
            custom_summary="New profile summary.",
            custom_experiences=experiences,
            candidate_profile={
                "name": "Jane Doe",
                "email": "jane@example.com",
                "linkedin": "linkedin.com/in/jane",
                "show_relocation_visa_line": True,
                "relocation_visa_line": "Eligible to work in EU.",
            },
        )

        out_doc = Document(self.output_path)
        self.assertEqual(out_doc.paragraphs[0].text.strip(), "Jane Doe")
        self.assertEqual(out_doc.paragraphs[1].text.strip(), "jane@example.com")
        self.assertIn("New profile summary.", out_doc.paragraphs[4].text)
        self.assertIn("Eligible to work in EU.", out_doc.paragraphs[4].text)

        out_table = out_doc.tables[0]
        self.assertEqual(out_table.cell(0, 0).text.strip(), "Manual Company")
        self.assertEqual(out_table.cell(0, 1).text.strip(), "Berlin, Germany")
        self.assertEqual(out_table.cell(1, 0).text.strip(), "Manual Role")
        self.assertEqual(out_table.cell(1, 1).text.strip(), "2024 - Present")

        bullet_lines = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
        headline_para = next((p for p in out_doc.paragraphs if p.text.strip() == "Impact headline"), None)
        self.assertIsNotNone(headline_para)
        self.assertNotIn("list", headline_para.style.name.lower())
        self.assertIsNotNone(headline_para.runs[0].font.size)
        self.assertAlmostEqual(headline_para.runs[0].font.size.pt, 10.0, places=1)
        self.assertIsNotNone(headline_para.paragraph_format.left_indent)
        self.assertAlmostEqual(headline_para.paragraph_format.left_indent.pt, 15.0, places=1)
        self.assertIn("Impact headline", bullet_lines)
        self.assertIn("New bullet one", bullet_lines)
        self.assertNotIn("- old bullet 2", bullet_lines)

    def test_structured_experience_appends_overflow_entries(self):
        legacy_key = self._build_template()
        company, title = _split_legacy_job_title(legacy_key)
        experiences = [
            {
                "anchor_key": "exp_1",
                "legacy_key": legacy_key,
                "company": company,
                "title": title,
                "date_range": "2022 - 2023",
                "location": "London, UK",
                "bullets": ["Base bullet"],
            },
            {
                "anchor_key": "exp_2",
                "company": "Overflow Co",
                "title": "Overflow Role",
                "date_range": "2024 - Present",
                "location": "Bristol, UK",
                "headline": "Overflow headline",
                "bullets": ["Overflow bullet"],
            },
        ]
        update_cv_bullets(
            input_file=self.template_path,
            output_file=self.output_path,
            custom_summary="New profile summary.",
            custom_experiences=experiences,
        )
        out_doc = Document(self.output_path)
        text_lines = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
        self.assertIn("Overflow Co | Overflow Role", text_lines)
        self.assertIn("2024 - Present | Bristol, UK", text_lines)
        self.assertIn("Overflow headline", text_lines)
        self.assertIn("• Overflow bullet", text_lines)


if __name__ == "__main__":
    unittest.main()
