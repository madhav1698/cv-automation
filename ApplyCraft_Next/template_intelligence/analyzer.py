import os
import re
from docx import Document
from typing import List, Dict
from .schemas import ParagraphData, StyleSignature, TemplateAnalysis, SectionInference, ExperienceBlock

class TemplateAnalyzer:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.paragraphs: List[ParagraphData] = []

    def analyze(self) -> TemplateAnalysis:
        """Core analysis pipeline"""
        self._extract_paragraphs()
        sections = self._infer_sections()
        blocks = self._infer_experience_blocks(sections)
        
        return TemplateAnalysis(
            paragraphs=self.paragraphs,
            inferred_sections=sections,
            inferred_experience_blocks=blocks
        )

    def _get_run_style(self, paragraph) -> StyleSignature:
        # Heuristic: use the style of the first run as representative
        if not paragraph.runs:
            return StyleSignature()
        
        first_run = paragraph.runs[0]
        # In python-docx, if a property is None, it inherits from style
        font = first_run.font
        
        return StyleSignature(
            style_name=paragraph.style.name,
            font_name=font.name or paragraph.style.font.name,
            font_size=font.size.pt if font.size else (paragraph.style.font.size.pt if paragraph.style.font.size else 12.0),
            is_bold=first_run.bold or paragraph.style.font.bold or False,
            is_italic=first_run.italic or paragraph.style.font.italic or False,
            alignment=str(paragraph.alignment or paragraph.style.paragraph_format.alignment or "left"),
            indent_left=paragraph.paragraph_format.left_indent.pt if paragraph.paragraph_format.left_indent else 0.0
        )

    def _extract_paragraphs(self):
        """Extracts content from body and tables in sequential order"""
        # Iterate through all top-level elements
        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                # It's a paragraph
                from docx.text.paragraph import Paragraph
                para = Paragraph(element, self.doc)
                self._process_para(para, location="body")
            elif element.tag.endswith('tbl'):
                # It's a table
                from docx.table import Table
                table = Table(element, self.doc)
                for row in table.rows:
                    for cell in row.cells:
                        # Process each paragraph inside the cell
                        for para in cell.paragraphs:
                            self._process_para(para, location="table_cell")

    def _process_para(self, para, location="body"):
        text = para.text.strip()
        if not text: return
        
        idx = len(self.paragraphs)
        sig = self._get_run_style(para)
        is_bullet = self._is_bullet_para(para)
        
        self.paragraphs.append(ParagraphData(
            index=idx,
            text=text,
            style_name=para.style.name,
            signature=sig,
            is_bullet=is_bullet,
            location=location
        ))

    def _is_bullet_para(self, para) -> bool:
        style_lower = para.style.name.lower()
        if "list" in style_lower or "bullet" in style_lower:
            return True
        # CHECK FOR SYMBOLS (even if style is Normal)
        text = para.text.strip()
        if text and text[0] in ["•", "-", "*", "▪", "▫", "➢", "→"]:
            return True
        return False

    def _infer_sections(self) -> List[SectionInference]:
        """Guesses section boundaries based on style signatures and keywords"""
        inferences = []
        experience_keywords = ["experience", "employment", "work history", "career"]
        
        for i, para in enumerate(self.paragraphs):
            text_lower = para.text.lower()
            
            # SECTION HEADING detection: Usually "Heading 1" or Bold+Uppercase+Short
            is_section_heading = "heading 1" in para.style_name.lower() or \
                                 "heading 2" in para.style_name.lower() or \
                                 (para.signature.is_bold and para.text.isupper() and len(para.text) < 40)
            
            if is_section_heading:
                if any(kw in text_lower for kw in experience_keywords):
                    # Guess this is the start of Experience
                    end_idx = self._find_next_section_heading(i + 1)
                    inferences.append(SectionInference(
                        section_type="Experience",
                        start_index=i,
                        end_index=end_idx,
                        confidence=0.9
                    ))
        return inferences

    def _find_next_section_heading(self, start_idx: int) -> int:
        for i in range(start_idx, len(self.paragraphs)):
            para = self.paragraphs[i]
            # Terminator is another section heading
            if "heading 1" in para.style_name.lower() or "heading 2" in para.style_name.lower() or \
               (para.signature.is_bold and para.text.isupper() and len(para.text) < 40):
                return i - 1
        return len(self.paragraphs) - 1

    def _infer_experience_blocks(self, sections: List[SectionInference]) -> List[ExperienceBlock]:
        """Deep dive into Experience section to find Company/Role/Date blocks"""
        blocks = []
        exp_section = next((s for s in sections if s.section_type == "Experience"), None)
        if not exp_section: return []
        
        # Stricter Date Regex
        date_range_pattern = re.compile(r"(\b(20\d{2}|19\d{2})\b\s*[\-–]\s*(Present|\b(20\d{2}|19\d{2})\b))", re.I)
        
        current_idx = exp_section.start_index + 1
        while current_idx <= exp_section.end_index:
            para = self.paragraphs[current_idx]
            
            # TRIGGER: A header line is NOT a bullet, is short, and NOT a full sentence
            is_header_line = self._is_header_likely(para)
            
            if is_header_line:
                block = ExperienceBlock()
                
                # HEADER GROUPING: Collect lines until we hit bullets or a "content paragraph"
                header_indices = []
                scan_idx = current_idx
                while scan_idx <= exp_section.end_index:
                    scan_para = self.paragraphs[scan_idx]
                    if scan_para.is_bullet or not self._is_header_likely(scan_para):
                        break
                    header_indices.append(scan_idx)
                    scan_idx += 1
                
                # Map the indices we found to fields with better heuristics
                location_keywords = ["uk", "usa", "sweden", "india", "germany", "london", "bristol", "stockholm"]
                
                for h_idx in header_indices:
                    h_para = self.paragraphs[h_idx]
                    txt = h_para.text.lower()
                    
                    if date_range_pattern.search(h_para.text):
                        block.date_idx = h_idx
                    elif any(kw in txt for kw in location_keywords) or "," in h_para.text:
                        # Likely a location
                        block.location_idx = h_idx
                    elif h_para.signature.is_italic and not block.role_idx:
                        # Role often italicized
                        block.role_idx = h_idx
                    elif h_para.signature.is_bold and not block.company_idx:
                        block.company_idx = h_idx
                    elif not block.role_idx:
                        block.role_idx = h_idx
                
                # Assign role if still missing
                if not block.role_idx and not block.company_idx and header_indices:
                    block.role_idx = header_indices[0]

                # FIND BULLET ZONE: All consecutive bullets OR descriptive paragraphs following this header group
                # Sometimes people use normal paragraphs for description before bullets
                search_idx = (max(header_indices) if header_indices else current_idx) + 1
                while search_idx <= exp_section.end_index:
                    p = self.paragraphs[search_idx]
                    if p.is_bullet:
                        if block.bullet_start_idx is None: block.bullet_start_idx = search_idx
                        block.bullet_end_idx = search_idx
                    elif self._is_header_likely(p):
                        # We hit the next job entry
                        break
                    else:
                        # Descriptive paragraph? Treat as part of bullet zone for now
                        if block.bullet_start_idx is None: block.bullet_start_idx = search_idx
                        block.bullet_end_idx = search_idx
                    search_idx += 1
                
                if block.company_idx or block.role_idx:
                    blocks.append(block)
                
                # Advance beyond everything we consumed
                current_idx = search_idx
            else:
                current_idx += 1
                
        return blocks

    def _is_header_likely(self, para: ParagraphData) -> bool:
        """Heuristic for a line being a header (Company/Role/Date) vs Content"""
        text = para.text.strip()
        if para.is_bullet: return False
        if len(text) < 3: return False
        
        # Headers are usually short
        if len(text) > 100: return False
        
        # Headers don't usually end with a period unless it's an abbreviation
        if text.endswith(".") and not re.search(r"\b(Inc|Ltd|Co|Corp)\.$", text, re.I):
            # If it's a long sentence with a period, it's likely content
            if len(text.split()) > 8: return False
            
        # Headers often have style traits
        has_style = para.signature.is_bold or para.signature.is_italic or para.style_name.startswith("Heading")
        
        # Or contain dates
        date_pattern = re.compile(r"\b(20\d{2}|19\d{2})\b|present", re.I)
        has_date = bool(date_pattern.search(text))
        
        return has_style or has_date or len(text) < 50
