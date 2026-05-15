from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from copy import deepcopy
from .schemas import TemplateConfig

class TemplateRenderer:
    def __init__(self, template_path: str, config: TemplateConfig):
        self.template_path = template_path
        self.config = config
        self.doc = Document(template_path)

    def render(self, data: dict, output_path: str):
        """
        Renders the tailored CV using learned anchors.
        Data: {"experience": [{"anchor_text": "Peermusic", "bullets": [...]}, ...]}
        """
        exp_data = data.get("experience", [])
        
        # 1. Update each Job's Bullets
        for entry in exp_data:
            self._update_job_bullets(entry)
            
        self.doc.save(output_path)

    def _update_job_bullets(self, entry):
        anchor_text = entry.get("anchor_text")
        new_bullets = entry.get("bullets", [])
        
        if not anchor_text: return
        
        # FIND THE ANCHOR in the current document
        # (This is why it's un-breakable: we search for the text)
        anchor_para_idx = self._find_para_by_text(anchor_text)
        if anchor_para_idx == -1: return
        
        # Find the "Bullet Zone" relative to anchor
        # For simplicity in V1: we look for existing bullets below the anchor
        # and replace them, or append if none found.
        
        # Let's find the first bullet paragraph after the anchor
        bullet_idx = -1
        for i in range(anchor_para_idx + 1, min(anchor_para_idx + 10, len(self.doc.paragraphs))):
            p = self.doc.paragraphs[i]
            if self._is_bullet(p):
                bullet_idx = i
                break
        
        if bullet_idx != -1:
            # We found existing bullets!
            # Replace the first one and delete the rest
            first_bullet_para = self.doc.paragraphs[bullet_idx]
            first_bullet_para.text = new_bullets[0] if new_bullets else ""
            
            # (In a more advanced version, we'd delete the old bullets and insert the new list properly)
            # For now, let's just update the text of existing ones
            for j, b_text in enumerate(new_bullets[1:]):
                # Simple append for demo
                new_p = self._insert_paragraph_after(self.doc.paragraphs[bullet_idx + j], b_text)
                new_p.style = first_bullet_para.style

    def _find_para_by_text(self, text):
        for i, p in enumerate(self.doc.paragraphs):
            if text in p.text:
                return i
        return -1

    def _is_bullet(self, para):
        style_lower = para.style.name.lower()
        if "list" in style_lower or "bullet" in style_lower: return True
        return para.text.strip().startswith(('•', '-', '*', '▪'))

    def _insert_paragraph_after(self, para, text):
        """Helper to insert a paragraph after another one in python-docx"""
        new_p = self.doc.add_paragraph(text)
        p = para._element
        p.addnext(new_p._element)
        return new_p
