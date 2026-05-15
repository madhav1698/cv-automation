import os
import copy
import shutil
import pythoncom
import comtypes.client
import PyPDF2
import sys
import re
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.paragraph import Paragraph

# Set up paths for internal imports
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(current_dir)
if project_root not in sys.path:
    sys.path.insert(0, project_root)

from helpers.logger import logger
from core.config import SUMMARY_TEXT, JOB_POSITIONS

dummy_bullets = [bullet for bullets in JOB_POSITIONS.values() for bullet in bullets]

def update_cv_bullets(input_file, output_file, custom_summary=None, custom_bullets=None, custom_headlines=None, current_location=None):
    """Read CV, update summary and bullet points, and save
    
    Args:
        input_file: Path to input CV template
        output_file: Path to save updated CV
        custom_summary: Optional custom summary text (uses default if None)
        custom_bullets: Optional dict mapping job titles to bullet lists, or flat list (legacy)
        custom_headlines: Optional dict mapping job titles to headline strings
        current_location: Optional current location (defaults to Stockholm)
    """
    doc = Document(input_file)
    
    # Use custom summary or default
    summary_text = custom_summary if custom_summary else SUMMARY_TEXT
    
    # Handle both dict (new job-aware) and list (legacy) formats
    if isinstance(custom_bullets, dict):
        bullets_dict = custom_bullets
        use_job_aware = True
    else:
        # Legacy: flat list
        bullets_list = custom_bullets if custom_bullets else dummy_bullets
        use_job_aware = False

    # First, replace any existing summary text that appears *before*
    # the fixed EU-citizen line, while keeping that line updated and BOLD.
    marker_base = "EU citizen, no visa required to work in EU."
    location_text = f"Currently in {current_location}, willing to relocate." if current_location else "Currently in Stockholm, willing to relocate."
    eu_marker_full = f"{marker_base} {location_text}"

    for paragraph in doc.paragraphs:
        if marker_base in paragraph.text:
            # Clear the paragraph and add summary text + bold EU line
            paragraph.clear()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run1 = paragraph.add_run(summary_text + " ")
            run1.font.name = 'Times New Roman'
            run1.font.size = Pt(10)
            run2 = paragraph.add_run(eu_marker_full)
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(10)
            run2.bold = True  # Make EU citizen line bold
            break
    
    def is_bullet_paragraph(paragraph):
        """Check if a paragraph is a bullet point"""
        # Check style name
        style_lower = paragraph.style.name.lower()
        if 'list' in style_lower or 'bullet' in style_lower:
            return True
        
        # Check paragraph format for list style
        if paragraph._element.pPr is not None:
            numPr = paragraph._element.pPr.numPr
            if numPr is not None:
                return True
        
        # Check text content for bullet characters
        text = paragraph.text.strip()
        bullet_chars = ['•', '-', '*', '▪', '▫', '○', '●', '→', '►']
        for char in bullet_chars:
            if text.startswith(char):
                return True
        
        return False
    
    def match_job_by_company(text, job_titles):
        """Match a job by company name (more flexible matching)"""
        text_upper = text.strip().upper()
        for job_title in job_titles:
            # Extract company name (part before – or -)
            job_parts = job_title.split("–") if "–" in job_title else job_title.split("-")
            if len(job_parts) >= 1:
                company_part = job_parts[0].strip().upper()
                company_words = company_part.replace(",", "").split()
                text_words = text_upper.replace(",", "").split()
                if company_words and company_words[0] in text_words:
                    return job_title
                if company_part in text_upper:
                    return job_title
        return None

    if use_job_aware:
        # Job-aware mode: collect bullets per job, then replace/add/remove as needed
        total_replaced = 0
        total_added = 0
        total_removed = 0
        
        # Build a map of table elements/paragraphs to their matched jobs
        table_to_job = {}
        for table in doc.tables:
            if len(table.rows) > 0:
                first_row = table.rows[0]
                for cell in first_row.cells:
                    cell_text = cell.text.strip()
                    matched_job = match_job_by_company(cell_text, bullets_dict.keys())
                    if matched_job:
                        table_to_job[table._element] = matched_job
                        break
        
        # First pass: collect all bullet paragraphs for each job
        job_to_bullets = {}  # Maps job title -> list of paragraph objects
        job_to_headline_para = {}
        job_to_header_element = {} 
        current_job = None
        
        paragraph_lookup = {para._element: para for para in doc.paragraphs}

        for element in doc.element.body:
            # Check if this is a table
            if element.tag.endswith('}tbl'):
                if element in table_to_job:
                    current_job = table_to_job[element]
                    job_to_bullets.setdefault(current_job, [])
                    job_to_header_element[current_job] = element
            
            # Check if this is a paragraph
            elif element.tag.endswith('}p'):
                para_obj = paragraph_lookup.get(element)
                if para_obj:
                    para_text = para_obj.text.strip()
                    matched_job = match_job_by_company(para_text, bullets_dict.keys())
                    if matched_job:
                        current_job = matched_job
                        job_to_bullets.setdefault(current_job, [])
                        job_to_header_element[current_job] = element
                        continue
                    
                    if current_job:
                        if is_bullet_paragraph(para_obj):
                            if para_obj.text.strip():
                                job_to_bullets[current_job].append(para_obj)
                        elif para_obj.text.strip() and not job_to_bullets.get(current_job) and current_job not in job_to_headline_para:
                            # It's a non-bullet paragraph before bullets.
                            # Skip if it looks like a role title or dates rather than a headline space
                            text_lower = para_obj.text.strip().lower()
                            
                            # Skip short date strings or location strings
                            if len(text_lower) < 30 and (
                                "202" in text_lower or 
                                "may" in text_lower or 
                                "sep" in text_lower or 
                                "bristol" in text_lower or 
                                "london" in text_lower or
                                "uk" in text_lower
                            ):
                                pass
                            # Skip known role-titles
                            elif any(role in text_lower for role in ["analyst", "scientist", "developer"]) and len(text_lower) < 40:
                                pass
                            else:
                                job_to_headline_para[current_job] = para_obj
        
        # Second pass: replace/add/remove bullets and headlines for each job
        
        # Fallback for templates without explicit headers (like Template 1)
        if len(job_to_bullets) == 0:
            logger.info("No headers matched. Falling back to sequential job bullet assignment.")
            all_bullet_paras = []
            for para in doc.paragraphs:
                if is_bullet_paragraph(para) and para.text.strip():
                    all_bullet_paras.append(para)
            
            job_to_bullets = {}
            current_bullet_idx = 0
            
            for job_title, original_bullets in JOB_POSITIONS.items():
                job_to_bullets.setdefault(job_title, [])
                num_to_take = len(original_bullets)
                end_idx = min(current_bullet_idx + num_to_take, len(all_bullet_paras))
                
                for i in range(current_bullet_idx, end_idx):
                    job_to_bullets[job_title].append(all_bullet_paras[i])
                    
                current_bullet_idx += num_to_take

        for job_title, new_bullets in bullets_dict.items():
            # Handle Headline
            new_headline = (custom_headlines or {}).get(job_title, "").strip()
            old_headline_para = job_to_headline_para.get(job_title)

            if new_headline:
                if old_headline_para:
                    old_headline_para.clear()
                    run = old_headline_para.add_run(new_headline)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.italic = True
                else:
                    header_element = job_to_header_element.get(job_title)
                    if header_element is not None:
                        parent = header_element.getparent()
                        insert_idx = list(parent).index(header_element) + 1
                        from docx.oxml import OxmlElement
                        new_p_element = OxmlElement('w:p')
                        parent.insert(insert_idx, new_p_element)
                        new_para = Paragraph(new_p_element, doc)
                        new_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        fmt = new_para.paragraph_format
                        fmt.space_before = Pt(2)
                        fmt.space_after = Pt(4)
                        fmt.left_indent = Pt(15) 
                        run = new_para.add_run(new_headline)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        run.font.italic = True
            elif old_headline_para:
                old_headline_para._element.getparent().remove(old_headline_para._element)

            # Handle Bullets
            if job_title not in job_to_bullets:
                logger.warning(f"Job '{job_title}' not found in document, skipping")
                continue
            
            old_bullet_paras = job_to_bullets[job_title]
            num_old = len(old_bullet_paras)
            num_new = len(new_bullets)
            
            for i in range(min(num_old, num_new)):
                para = old_bullet_paras[i]
                para.clear()
                run = para.add_run(new_bullets[i])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                total_replaced += 1
            
            if num_new > num_old and old_bullet_paras:
                last_para = old_bullet_paras[-1]
                last_element = last_para._element
                parent = last_element.getparent()
                insert_index = list(parent).index(last_element) + 1
                for i in range(num_old, num_new):
                    new_p_element = copy.deepcopy(last_element)
                    for run_elem in new_p_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        run_elem.text = ''
                    parent.insert(insert_index, new_p_element)
                    insert_index += 1
                    new_para = Paragraph(new_p_element, last_para._parent)
                    new_para.clear()
                    run = new_para.add_run(new_bullets[i])
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    total_added += 1
            elif num_new < num_old:
                for i in range(num_new, num_old):
                    para = old_bullet_paras[i]
                    para._element.getparent().remove(para._element)
                    total_removed += 1
        
        logger.info(f"Updated CV saved to: {output_file}")
        logger.info(f"Replaced: {total_replaced}, Added: {total_added}, Removed: {total_removed} bullet points")
        
    else:
        # Legacy mode: replace all bullets sequentially
        bullet_index = 0
        
        def update_paragraph_legacy(paragraph):
            """Update a single paragraph with bullet text and set font size to 10"""
            nonlocal bullet_index
            if is_bullet_paragraph(paragraph) and paragraph.text.strip():
                text = bullets_list[bullet_index % len(bullets_list)]
                paragraph.clear()
                run = paragraph.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                bullet_index += 1
                return True
            return False
        
        # Iterate through all paragraphs in the document
        for paragraph in doc.paragraphs:
            update_paragraph_legacy(paragraph)
        
        # Also check tables for bullet points
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        update_paragraph_legacy(paragraph)
        
        logger.info(f"Updated CV saved to: {output_file}")
        logger.info(f"Replaced {bullet_index} bullet points with dummy text")
    
    # Remove trailing empty paragraphs to prevent blank pages
    logger.debug("Cleaning up trailing empty paragraphs...")
    while len(doc.paragraphs) > 0:
        last_para = doc.paragraphs[-1]
        if not last_para.text.strip():
            # Check if it has any content like images/tables (unlikely in para, but good to be safe)
            # For now, just check text.
            p_element = last_para._element
            if p_element.getparent() is not None:
                p_element.getparent().remove(p_element)
                # doc.paragraphs re-evaluates, so loop continues with new last para
        else:
            break
            
    # Remove any trailing section breaks that might cause a blank page
    if len(doc.sections) > 1:
        # Check the LAST paragraph for a section break (sectPr)
        # This is what usually causes the 'blank page' in Template 2
        last_para = doc.paragraphs[-1] if doc.paragraphs else None
        if last_para and last_para._p.pPr is not None and last_para._p.pPr.sectPr is not None:
            logger.debug("Removing trailing section break...")
            
            # Capture properties of the 'good' section (usually the first one)
            first_sec = doc.sections[0]
            good_margins = {
                'bottom': first_sec.bottom_margin,
                'top': first_sec.top_margin,
                'left': first_sec.left_margin,
                'right': first_sec.right_margin
            }
            
            # Remove the break
            last_para._p.pPr.remove(last_para._p.pPr.sectPr)
            
            # After removal, the remaining last section might have inherited 
            # the 'bad' margins from the former last section. Restore them.
            new_last_sec = doc.sections[-1]
            new_last_sec.bottom_margin = good_margins['bottom']
            new_last_sec.top_margin = good_margins['top']
            new_last_sec.left_margin = good_margins['left']
            new_last_sec.right_margin = good_margins['right']

    # Save the updated document
    doc.save(output_file)
    return output_file
def convert_to_pdf(docx_file, pdf_file):
    """Convert DOCX to PDF"""
    # Initialize COM for this thread
    pythoncom.CoInitialize()
    
    try:
        from docx2pdf import convert
        convert(docx_file, pdf_file)
        logger.info(f"PDF created: {pdf_file}")
        remove_blank_pages(pdf_file)
        return True
    except Exception as e:
        logger.warning(f"docx2pdf conversion failed: {e}")
        logger.info("Trying alternative method with comtypes...")
        try:
            # Ensure paths are absolute
            doc_path = os.path.abspath(docx_file)
            pdf_path = os.path.abspath(pdf_file)
            
            # Initialize Word
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            
            try:
                doc = word.Documents.Open(doc_path)
                # 17 = wdExportFormatPDF
                doc.SaveAs(pdf_path, FileFormat=17)
                doc.Close()
                logger.info(f"PDF created via comtypes: {pdf_file}")
                remove_blank_pages(pdf_path)
                return True
            finally:
                word.Quit()
                
        except Exception as e2:
            logger.error(f"Alternative conversion also failed: {e2}")
            return False
    finally:
        # Uninitialize COM for this thread
        pythoncom.CoUninitialize()

def remove_blank_pages(pdf_path):
    """Remove completely blank pages from a PDF file."""
    try:
        if not os.path.exists(pdf_path):
            return

        # Create a temporary output file
        temp_output = pdf_path.replace(".pdf", "_temp.pdf")
        
        has_changes = False
        original_page_count = 0
        new_page_count = 0

        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            writer = PyPDF2.PdfWriter()
            
            original_page_count = len(reader.pages)
            
            for i, page in enumerate(reader.pages):
                # Check for text content - more robust check
                text = page.extract_text() or ""
                cleaned_text = text.strip()
                # Consider page blank if text is very short (e.g. just page numbers)
                has_significant_text = len(cleaned_text) > 20 
                
                # Check for images (if available in this PyPDF2 version)
                has_images = False
                try:
                    if hasattr(page, 'images') and page.images:
                        # Only count substantial images (hard to tell size, but existence usually enough)
                        has_images = True
                except:
                    # If image extraction fails, assume no images or ignore error
                    pass
                
                # Setup page content check - strict "completely blank" or just artifacts
                if has_significant_text or has_images:
                    writer.add_page(page)
                    new_page_count += 1
                else:
                    logger.debug(f"Removing blank/sparse page {i+1} from PDF (text len: {len(cleaned_text)})")
                    has_changes = True

            # If no changes, just return (unless we want to rebuild for other reasons)
            if not has_changes:
                return

            with open(temp_output, 'wb') as out_f:
                writer.write(out_f)

        # Replace original file with cleaned version
        shutil.move(temp_output, pdf_path)
        logger.info(f"PDF cleaned: {original_page_count} -> {new_page_count} pages")
        
    except Exception as e:
        logger.error(f"Error removing blank pages: {e}")
        # Clean up temp file if it exists
        if os.path.exists(temp_output):
            try:
                os.remove(temp_output)
            except:
                pass

if __name__ == "__main__":
    import sys
    
    input_file = "Madhav_Manohar Gopal_CV.docx"
    
    # Get company name from command line argument or prompt user
    company_name = None
    if len(sys.argv) > 1:
        company_name = sys.argv[1].strip()
    
    if not company_name:
        company_name = input("Enter company name (or press Enter to skip): ").strip()
    
    if not company_name:
        company_name = "Updated"  # Default fallback
    
    # Clean company name for filename (remove special characters, replace spaces with underscores)
    company_name_clean = company_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
    company_name_clean = "".join(c for c in company_name_clean if c.isalnum() or c in ("_", "-"))
    
    # Create outputs folder structure with date-based subfolders
    from datetime import datetime
    base_dir = os.path.dirname(input_file) or "."
    outputs_dir = os.path.join(base_dir, "outputs")
    
    # Create date folder name (format: D-M-YY)
    today = datetime.now()
    date_folder_name = f"{today.day}-{today.month}-{today.strftime('%y')}"  # e.g., "3-2-26"
    date_output_dir = os.path.join(outputs_dir, date_folder_name)
    
    # Create company folder inside date folder
    company_output_dir = os.path.join(date_output_dir, company_name_clean)
    os.makedirs(company_output_dir, exist_ok=True)
    
    # Generate output filenames in company folder
    output_docx = os.path.join(company_output_dir, f"Madhav_Manohar_Gopal_CV_{company_name_clean}.docx")
    output_pdf = os.path.join(company_output_dir, f"Madhav_Manohar_Gopal_CV_{company_name_clean}.pdf")
    
    print(f"\nGenerating CV for: {company_name}")
    print(f"Output files: {output_docx} and {output_pdf}\n")

    # Update bullets
    update_cv_bullets(input_file, output_docx)

    # Convert to PDF
    convert_to_pdf(output_docx, output_pdf)
