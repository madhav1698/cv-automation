import os
import sys
from docx import Document
from docx.shared import Pt
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set up paths for internal imports
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(current_dir)
if project_root not in sys.path:
    sys.path.insert(0, project_root)

from helpers.logger import logger

def generate_cover_letter(output_path, company_name, city, country, date_str, body_text, hiring_manager="Hiring Manager", candidate_name="Madhav Manohar Gopal"):
    """
    Generates a cover letter docx file with premium formatting.
    """
    # Only title case if it's a short string (likely a name), not if it's the whole body by mistake
    if len(hiring_manager) < 100:
        hiring_manager = hiring_manager.title()
    
    doc = Document()
    
    # Set Margins (Standard 1 inch)
    from docx.shared import Inches
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set paragraph spacing to zero for the whole document
    # We will control spacing by adding empty paragraphs
    for s in doc.styles:
        if hasattr(s, 'paragraph_format'):
            s.paragraph_format.space_before = Pt(0)
            s.paragraph_format.space_after = Pt(0)
            s.paragraph_format.line_spacing = 1.0

    def add_line(text="", is_bold=False, alignment=None):
        p = doc.add_paragraph()
        if alignment:
            p.alignment = alignment
        if text:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = is_bold
            # Force font for compatibility
            from docx.oxml.ns import qn
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        return p

    # Add Header (Always 'Hiring Manager' for the title)
    add_line("Hiring Manager")
    add_line(company_name)
    add_line(f"{city}, {country}")
    add_line() # Space
    add_line(date_str)
    add_line() # Space
    
    # Add Greeting (Uses specific name if provided)
    add_line(f"Dear {hiring_manager},")
    add_line() # Space
    
    # Add Body
    body_paragraphs = [p.strip() for p in body_text.split('\n') if p.strip()]
    for i, p_text in enumerate(body_paragraphs):
        add_line(p_text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        add_line() # Space after every paragraph
    
    # Add Closing
    add_line("Sincerely,")
    add_line(candidate_name)
    
    # Save
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    # Test
    test_body = (
        "At scale, employee listening fails for one simple reason: feedback is collected faster than organisations "
        "can decide what to do with it. What drew me to this role at Proofpoint is that it treats listening as a system "
        "with ownership, governance, and consequences, not just a survey cycle.\n\n"
        "In my work, the most difficult part has never been analysis. It has been deciding which signals deserve attention."
    )
    generate_cover_letter(
        "test_cover_letter.docx",
        "Proofpoint",
        "Cork",
        "Ireland",
        "06/02/2026",
        test_body
    )
    print("Test cover letter generated.")
